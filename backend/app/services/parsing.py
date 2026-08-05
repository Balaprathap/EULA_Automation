"""Document parsing and upload validation.

Handles text-based PDF, DOCX, and TXT. Validation is defence-in-depth: the file
is sniffed by magic bytes first (the extension is only a secondary signal),
size and page budgets are enforced before any expensive work, and scanned or
encrypted PDFs are rejected with an explicit, actionable message rather than
being silently analyzed as empty documents.
"""

from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass, field
from typing import Any

from app.core.errors import (
    EmptyDocument,
    EncryptedDocument,
    FileTooLarge,
    ScannedDocument,
    Unprocessable,
    UnsupportedMediaType,
)
from app.services.normalization import content_hash, normalize_text, text_density

SUPPORTED_MIME = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "text/plain": "txt",
}
SUPPORTED_EXTENSIONS = {".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".text": "txt"}

# A text PDF averages thousands of characters per page. Below this, the pages
# are almost certainly images.
MIN_CHARS_PER_PAGE = 120
MIN_TOTAL_CHARS = 200

# Guards against zip bombs in DOCX containers.
MAX_DECOMPRESSION_RATIO = 200
MAX_DECOMPRESSED_BYTES = 400 * 1024 * 1024


@dataclass
class ParsedDocument:
    normalized_text: str
    page_count: int
    char_count: int
    source_type: str
    content_sha256: str
    title: str | None = None
    metadata: dict[str, object] = field(default_factory=dict)


def sniff_file_type(data: bytes, filename: str = "") -> str:
    """Identify the file type from magic bytes; the extension is a fallback only."""
    if not data:
        raise EmptyDocument("The uploaded file is empty.")

    if data[:5] == b"%PDF-":
        return "pdf"

    # DOCX is a ZIP container; confirm the Word content type is actually present.
    if data[:4] in (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08"):
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                names = set(archive.namelist())
                if "word/document.xml" in names:
                    return "docx"
        except zipfile.BadZipFile:
            pass
        raise UnsupportedMediaType(
            "This ZIP-based file is not a Word (.docx) document.",
        )

    # Plain text: must decode as UTF-8/Latin-1 and contain no NUL bytes.
    if b"\x00" not in data[:8192]:
        try:
            data[:8192].decode("utf-8")
            return "txt"
        except UnicodeDecodeError:
            pass

    extension = ""
    if "." in filename:
        extension = filename[filename.rindex(".") :].lower()
    if extension in SUPPORTED_EXTENSIONS:
        raise UnsupportedMediaType(
            f"The file extension suggests {SUPPORTED_EXTENSIONS[extension].upper()}, "
            "but the file contents do not match that format."
        )
    raise UnsupportedMediaType()


def validate_size(data: bytes, max_bytes: int) -> None:
    if len(data) > max_bytes:
        raise FileTooLarge(
            f"The file is {len(data) / 1_048_576:.1f} MB, which exceeds the "
            f"{max_bytes / 1_048_576:.0f} MB limit."
        )
    if not data:
        raise EmptyDocument("The uploaded file is empty.")


def parse_pdf(data: bytes, max_pages: int) -> ParsedDocument:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(io.BytesIO(data))
    except PdfReadError as exc:
        raise Unprocessable(f"This PDF could not be read: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        # An empty user password is common and harmless; try it before rejecting.
        try:
            if reader.decrypt("") == 0:
                raise EncryptedDocument()
        except Exception as exc:  # noqa: BLE001 - any decrypt failure means unreadable
            raise EncryptedDocument() from exc

    page_count = len(reader.pages)
    if page_count == 0:
        raise EmptyDocument("This PDF contains no pages.")
    if page_count > max_pages:
        raise Unprocessable(
            f"This document has {page_count} pages, which exceeds the {max_pages}-page limit.",
            code="TOO_MANY_PAGES",
        )

    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - a single bad page must not fail the upload
            parts.append("")

    raw = "\n\n".join(parts)
    normalized = normalize_text(raw)

    # Scanned-PDF detection: reject explicitly rather than analyzing nothing.
    if (
        len(normalized) < MIN_TOTAL_CHARS
        or text_density(normalized, page_count) < MIN_CHARS_PER_PAGE
    ):
        raise ScannedDocument()

    info: Any = reader.metadata or {}
    title = None
    raw_title = info.get("/Title") if hasattr(info, "get") else None
    if raw_title and str(raw_title).strip():
        title = str(raw_title).strip()[:300]

    return ParsedDocument(
        normalized_text=normalized,
        page_count=page_count,
        char_count=len(normalized),
        source_type="pdf",
        content_sha256=content_hash(normalized),
        title=title,
        metadata={"pdf_version": data[5:8].decode("ascii", "ignore")},
    )


def parse_docx(data: bytes, max_pages: int) -> ParsedDocument:
    import docx

    # Zip-bomb guard before handing the archive to python-docx.
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            compressed = sum(i.compress_size for i in archive.infolist()) or 1
            decompressed = sum(i.file_size for i in archive.infolist())
        if (
            decompressed > MAX_DECOMPRESSED_BYTES
            or decompressed / compressed > MAX_DECOMPRESSION_RATIO
        ):
            raise Unprocessable(
                "This file expands to an unsafe size and was rejected.",
                code="SUSPICIOUS_ARCHIVE",
            )
    except zipfile.BadZipFile as exc:
        raise UnsupportedMediaType("This file is not a valid .docx document.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise Unprocessable("This Word document could not be read.") from exc

    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    normalized = normalize_text("\n".join(parts))
    if len(normalized) < MIN_TOTAL_CHARS:
        raise EmptyDocument("This Word document contains no readable text.")

    # DOCX has no fixed pagination; estimate at ~3000 characters per page.
    page_count = max(1, len(normalized) // 3000)
    if page_count > max_pages:
        raise Unprocessable(
            f"This document is approximately {page_count} pages, which exceeds "
            f"the {max_pages}-page limit.",
            code="TOO_MANY_PAGES",
        )

    title = None
    for paragraph in document.paragraphs[:10]:
        if paragraph.text.strip():
            title = paragraph.text.strip()[:300]
            break

    return ParsedDocument(
        normalized_text=normalized,
        page_count=page_count,
        char_count=len(normalized),
        source_type="docx",
        content_sha256=content_hash(normalized),
        title=title,
        metadata={"paragraph_count": len(document.paragraphs)},
    )


def parse_txt(data: bytes, max_pages: int) -> ParsedDocument:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            raw = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise UnsupportedMediaType("This text file uses an unsupported character encoding.")

    return parse_plain_text(raw, max_pages, source_type="txt")


def parse_plain_text(raw: str, max_pages: int, source_type: str = "paste") -> ParsedDocument:
    """Parse pasted text. Same normalization path as uploaded files."""
    normalized = normalize_text(raw)
    if len(normalized) < MIN_TOTAL_CHARS:
        raise EmptyDocument(
            f"The agreement text is too short to analyze "
            f"({len(normalized)} characters; at least {MIN_TOTAL_CHARS} are required)."
        )

    page_count = max(1, len(normalized) // 3000)
    if page_count > max_pages:
        raise Unprocessable(
            f"This text is approximately {page_count} pages, which exceeds "
            f"the {max_pages}-page limit.",
            code="TOO_MANY_PAGES",
        )

    title = None
    for line in normalized.split("\n")[:5]:
        candidate = line.strip()
        if 8 <= len(candidate) <= 200:
            title = candidate
            break

    return ParsedDocument(
        normalized_text=normalized,
        page_count=page_count,
        char_count=len(normalized),
        source_type=source_type,
        content_sha256=content_hash(normalized),
        title=title,
    )


def parse_upload(data: bytes, filename: str, *, max_bytes: int, max_pages: int) -> ParsedDocument:
    """Validate and parse an uploaded file. The single entry point for uploads."""
    validate_size(data, max_bytes)
    file_type = sniff_file_type(data, filename)

    if file_type == "pdf":
        parsed = parse_pdf(data, max_pages)
    elif file_type == "docx":
        parsed = parse_docx(data, max_pages)
    else:
        parsed = parse_txt(data, max_pages)

    if not parsed.title and filename:
        parsed.title = re.sub(r"\.[^.]+$", "", filename)[:300] or None
    return parsed
